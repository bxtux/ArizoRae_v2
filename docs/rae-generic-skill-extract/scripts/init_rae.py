#!/usr/bin/env python3
"""
init_rae.py
Outillage pour le workflow /init du skill rae-generic.

Rôle : partie déterministe du setup initial. Agnostique au métier :
aucune liste de professions, aucun catalogue figé. Le preset métier
est produit hors script (recherche web pilotée par l'agent), puis
cette partie ne fait que mémoriser son emplacement.

    1. Extraire le texte d'un CV (PDF, DOCX, ODT, TXT, MD).
    2. Poser la configuration (emplacement des fichiers de profil,
       métier déclaré en texte libre, pays).
    3. Matérialiser les squelettes FACTS.md et BULLET_LIBRARY.md à
       partir des templates si les fichiers cibles n'existent pas.
    4. Écrire user/config.json dans le skill, qui mémorise les
       chemins et les choix.

Le parsing sémantique proprement dit (qui dit quoi dans le CV) est
laissé à l'agent. Ce script prépare le terrain et fournit le texte
brut.

Dépendances optionnelles détectées à l'exécution :
    - pdfplumber ou pypdf (PDF)
    - python-docx (DOCX)
    - odfpy (ODT)

Usage typique :

    python init_rae.py extract --cv path/to/cv.pdf \\
        --out /home/user/.rae/cv_raw.txt
    python init_rae.py configure \\
        --output-dir /home/user/.rae \\
        --metier "administrateur systèmes Linux" \\
        --country BE \\
        --lang fr
    python init_rae.py status
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from typing import Optional


SKILL_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = SKILL_ROOT / "templates"
USER_DIR = SKILL_ROOT / "user"
CONFIG_PATH = USER_DIR / "config.json"


# ---------------------------------------------------------------------------
# Extraction de texte
# ---------------------------------------------------------------------------


def extract_pdf(path: Path) -> str:
    """Retourne le texte brut d'un PDF. Essaye pdfplumber, puis pypdf,
    puis pdftotext en dernier recours."""
    try:
        import pdfplumber  # type: ignore
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                pages.append(txt)
        return "\n\n".join(pages).strip()
    except ImportError:
        pass

    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        return "\n\n".join((p.extract_text() or "") for p in reader.pages).strip()
    except ImportError:
        pass

    pdftotext = shutil.which("pdftotext")
    if pdftotext:
        proc = subprocess.run(
            [pdftotext, "-layout", str(path), "-"],
            capture_output=True,
            text=True,
            check=False,
        )
        if proc.returncode == 0:
            return proc.stdout.strip()

    raise RuntimeError(
        "Aucun extracteur PDF disponible. Installe pdfplumber, pypdf, "
        "ou le binaire poppler (pdftotext)."
    )


def extract_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "python-docx non installé. pip install python-docx."
        ) from exc

    doc = Document(str(path))
    lines = [para.text for para in doc.paragraphs]
    # Tables aussi, souvent utilisées dans les CV
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(l for l in lines if l.strip()).strip()


def extract_odt(path: Path) -> str:
    """Extraction minimaliste via le XML interne. Évite une dépendance
    lourde à odfpy si elle n'est pas déjà là."""
    try:
        from odf.opendocument import load  # type: ignore
        from odf.text import P  # type: ignore

        doc = load(str(path))
        paragraphs = doc.getElementsByType(P)
        parts = []
        for p in paragraphs:
            parts.append("".join(node.data for node in p.childNodes
                                 if getattr(node, "data", None)))
        return "\n".join(parts).strip()
    except ImportError:
        pass

    # Fallback brut : content.xml lu et nettoyé grossièrement
    import re
    with zipfile.ZipFile(path) as zf:
        with zf.open("content.xml") as fh:
            raw = fh.read().decode("utf-8", errors="replace")
    text = re.sub(r"<[^>]+>", "\n", raw)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def extract_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


def extract_cv(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".pdf":
        return extract_pdf(path)
    if suf == ".docx":
        return extract_docx(path)
    if suf == ".odt":
        return extract_odt(path)
    if suf in {".txt", ".md", ".markdown"}:
        return extract_plain(path)
    raise ValueError(f"Format non supporté pour le CV : {suf}")


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------


def load_config() -> dict:
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass
    return {}


def save_config(cfg: dict) -> None:
    USER_DIR.mkdir(parents=True, exist_ok=True)
    CONFIG_PATH.write_text(
        json.dumps(cfg, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def resolve_output_dir(choice: Optional[str]) -> Path:
    """Résout le répertoire de stockage FACTS/BULLET selon le choix.
    Codes reconnus : 'local' (dans le skill), 'home' (~/.rae),
    ou chemin absolu/relatif."""
    if choice in (None, "", "local"):
        return USER_DIR
    if choice == "home":
        return Path.home() / ".rae"
    return Path(choice).expanduser().resolve()


def ensure_profile_files(output_dir: Path) -> tuple[Path, Path, Path]:
    """Crée output_dir si besoin et matérialise FACTS.md, BULLET_LIBRARY.md
    et un preset.md vide (à remplir par l'agent via recherche web)."""
    output_dir.mkdir(parents=True, exist_ok=True)
    facts = output_dir / "FACTS.md"
    bullet = output_dir / "BULLET_LIBRARY.md"
    preset = output_dir / "preset.md"

    if not facts.exists():
        tpl = TEMPLATES_DIR / "FACTS.template.md"
        if tpl.exists():
            facts.write_text(tpl.read_text(encoding="utf-8"), encoding="utf-8")
        else:
            facts.write_text("# FACTS\n\n(à compléter)\n", encoding="utf-8")

    if not bullet.exists():
        tpl = TEMPLATES_DIR / "BULLET_LIBRARY.template.md"
        if tpl.exists():
            bullet.write_text(tpl.read_text(encoding="utf-8"), encoding="utf-8")
        else:
            bullet.write_text("# BULLET LIBRARY\n\n(à compléter)\n", encoding="utf-8")

    if not preset.exists():
        preset.write_text(
            "# Preset métier\n\n"
            "Fichier en attente de génération.\n\n"
            "L'agent doit le remplir via recherche web lors de /init, "
            "en suivant la trame indiquée dans SKILL.md (sections Périmètre, "
            "Thèmes de bullets, Différenciateurs, Jobboards, Questions "
            "d'entretien, Pièges, Mots-clés ATS, Sources).\n",
            encoding="utf-8",
        )

    return facts, bullet, preset


# ---------------------------------------------------------------------------
# Commandes CLI
# ---------------------------------------------------------------------------


def cmd_extract(args) -> int:
    path = Path(args.cv).expanduser().resolve()
    if not path.exists():
        print(f"[erreur] CV introuvable : {path}", file=sys.stderr)
        return 2
    try:
        text = extract_cv(path)
    except Exception as exc:
        print(f"[erreur] extraction : {exc}", file=sys.stderr)
        return 3

    if args.out:
        out_path = Path(args.out).expanduser().resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(text, encoding="utf-8")
        print(f"[ok] texte CV écrit dans {out_path} ({len(text)} caractères)")
    else:
        sys.stdout.write(text)
        sys.stdout.write("\n")
    return 0


def cmd_configure(args) -> int:
    cfg = load_config()

    output_dir = resolve_output_dir(args.output_dir)
    facts, bullet, preset = ensure_profile_files(output_dir)

    # Le preset_status vaut "pending" tant que l'agent n'a pas
    # effectivement rempli preset.md via recherche web.
    preset_status = cfg.get("preset_status", "pending")
    if args.preset_status:
        preset_status = args.preset_status

    cfg.update({
        "output_dir": str(output_dir),
        "facts_path": str(facts),
        "bullet_path": str(bullet),
        "preset_path": str(preset),
        "preset_status": preset_status,
        "metier": args.metier or cfg.get("metier", ""),
        "country": args.country or cfg.get("country", "international"),
        "lang": args.lang or cfg.get("lang", "auto"),
    })
    if args.profile_name:
        cfg["profile_name"] = args.profile_name

    save_config(cfg)

    print(f"[ok] configuration enregistrée dans {CONFIG_PATH}")
    print(json.dumps(cfg, indent=2, ensure_ascii=False))
    if preset_status == "pending":
        print(
            "\n[rappel] preset.md reste à générer par recherche web. "
            "Voir la section 'Construction du preset via recherche web' "
            "dans SKILL.md."
        )
    return 0


def cmd_status(args) -> int:
    cfg = load_config()
    if not cfg:
        print("[info] aucun profil initialisé. Lance `configure` puis renseigne "
              "FACTS.md et BULLET_LIBRARY.md.")
        return 0

    print(json.dumps(cfg, indent=2, ensure_ascii=False))
    facts = Path(cfg.get("facts_path", "")).expanduser()
    bullet = Path(cfg.get("bullet_path", "")).expanduser()
    preset = Path(cfg.get("preset_path", "")).expanduser()
    print()
    print(f"FACTS.md        : {'présent' if facts.exists() else 'manquant'} ({facts})")
    print(f"BULLET_LIBRARY  : {'présent' if bullet.exists() else 'manquant'} ({bullet})")
    preset_state = cfg.get("preset_status", "pending")
    marker = "présent" if preset.exists() else "manquant"
    print(f"preset.md       : {marker} / statut={preset_state} ({preset})")
    return 0


def cmd_reset(args) -> int:
    if CONFIG_PATH.exists():
        CONFIG_PATH.unlink()
        print(f"[ok] {CONFIG_PATH} supprimé.")
    else:
        print("[info] rien à supprimer.")
    return 0


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="init_rae",
        description="Initialisation du profil pour le skill rae-generic.",
    )
    sub = p.add_subparsers(dest="command", required=True)

    pe = sub.add_parser("extract", help="Extraire le texte d'un CV.")
    pe.add_argument("--cv", required=True, help="Chemin vers le CV.")
    pe.add_argument("--out", help="Fichier de sortie. Par défaut, stdout.")
    pe.set_defaults(func=cmd_extract)

    pc = sub.add_parser("configure", help="Écrire user/config.json du skill.")
    pc.add_argument(
        "--output-dir",
        help="Emplacement FACTS/BULLET/preset. 'local', 'home', ou chemin.",
    )
    pc.add_argument(
        "--metier",
        help=(
            "Métier déclaré en texte libre. Aucun catalogue, aucune "
            "contrainte de vocabulaire. Le preset sera construit par "
            "recherche web."
        ),
    )
    pc.add_argument(
        "--preset-status",
        choices=("pending", "ready", "incomplete"),
        help=(
            "Statut du preset.md. 'pending' par défaut à l'init, "
            "'ready' après génération par l'agent, 'incomplete' si "
            "le web était inaccessible."
        ),
    )
    pc.add_argument("--country", help="Code pays (BE, FR, CH, CA, international...).")
    pc.add_argument("--lang", help="Langue par défaut des documents (fr, en, auto).")
    pc.add_argument("--profile-name", help="Étiquette libre pour ce profil.")
    pc.set_defaults(func=cmd_configure)

    ps = sub.add_parser("status", help="Afficher la configuration courante.")
    ps.set_defaults(func=cmd_status)

    pr = sub.add_parser("reset", help="Supprimer la configuration.")
    pr.set_defaults(func=cmd_reset)

    return p


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
